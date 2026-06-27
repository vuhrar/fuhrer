# doc_processing.py
"""
استخراج النص من ملفات PDF/DOCX/TXT/HTML/CSV و
دوال استخراج مواد قانونية من النصوص.
مع ترقية لاستخراج الكيانات العمالية (الراتب، المدة، تاريخ الفصل).
"""
from typing import List, Dict, Any
import io, re
from datetime import datetime
from utils import _bytes, _norm

# ==========================
# دوال استخراج الكيانات العمالية
# ==========================

def extract_labor_entities(text: str) -> Dict[str, Any]:
    """
    استخراج الكيانات العمالية من النص:
    - الراتب الأساسي
    - الراتب الإجمالي
    - مدة الخدمة (بالسنوات)
    - تاريخ الفصل
    - نوع الفصل (تعسفي/طوعي)
    """
    entities = {
        "basic_salary": None,
        "total_salary": None,
        "service_years": None,
        "termination_date": None,
        "termination_type": None,  # "arbitrary", "voluntary", "unknown"
        "absence_days": None,
        "salary_delay_months": None,
    }

    # استخراج الراتب الأساسي
    basic_match = re.search(r'(?:الراتب\s+الأساسي|الأساسي)\s*[:]\s*([\d,]+)', text, re.IGNORECASE)
    if basic_match:
        entities["basic_salary"] = float(basic_match.group(1).replace(',', ''))

    # استخراج الراتب الإجمالي
    total_match = re.search(r'(?:الراتب\s+الإجمالي|الإجمالي)\s*[:]\s*([\d,]+)', text, re.IGNORECASE)
    if total_match:
        entities["total_salary"] = float(total_match.group(1).replace(',', ''))

    # استخراج مدة الخدمة (بالسنوات)
    years_match = re.search(r'(?:مدة\s+الخدمة|الخدمة)\s*[:]\s*([\d,.]+)\s*(?:سنوات|سنة)', text, re.IGNORECASE)
    if years_match:
        entities["service_years"] = float(years_match.group(1).replace(',', ''))

    # استخراج تاريخ الفصل
    date_match = re.search(r'(?:تاريخ\s+الفصل|الفصل)\s*[:]\s*([\d/]+)', text, re.IGNORECASE)
    if date_match:
        entities["termination_date"] = date_match.group(1)

    # تحديد نوع الفصل
    if "تعسفي" in text or "غير مشروع" in text:
        entities["termination_type"] = "arbitrary"
    elif "استقالة" in text or "طوعي" in text:
        entities["termination_type"] = "voluntary"
    else:
        entities["termination_type"] = "unknown"

    # أيام الغياب
    absence_match = re.search(r'(?:أيام\s+الغياب|الغياب)\s*[:]\s*([\d]+)', text, re.IGNORECASE)
    if absence_match:
        entities["absence_days"] = int(absence_match.group(1))

    # تأخير الراتب (شهور)
    delay_match = re.search(r'(?:تأخير\s+الراتب|متأخرات)\s*[:]\s*([\d]+)\s*(?:شهر|شهور)', text, re.IGNORECASE)
    if delay_match:
        entities["salary_delay_months"] = int(delay_match.group(1))

    return entities

# ==========================
# دوال استخراج المواد القانونية (محسّنة)
# ==========================

def extract_laws_from_text(text: str, source: str = "") -> List[Dict]:
    records = []
    lines = text.split('\n')
    current_article = ""
    current_law = "الأنظمة السعودية"
    capture = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # الكشف عن اسم النظام (يحدث مرة واحدة)
        law_match = re.search(r'نظام\s+[\u0600-\u06ff\s]{4,35}(?=\n|،|\.)', line)
        if law_match and not any(r.get("law_name") == law_match.group(0) for r in records[-5:]):
            current_law = law_match.group(0).strip()

        # الكشف الدقيق عن بداية المادة (يدعم الأرقام العربية والهندية)
        article_match = re.match(r'(المادة|مادة)\s+([\u0660-\u0669\d]+)\s*[:\.\-]', line)
        if article_match:
            if capture and len(current_article) > 50:
                clean = re.sub(r'\s+', ' ', current_article).strip()
                if sum(1 for c in clean if '\u0600' <= c <= '\u06ff') > 15:
                    records.append({
                        "text": clean[:1200],
                        "article": f"{article_match.group(1)} {article_match.group(2)}",
                        "law_name": current_law,
                        "source": source,
                        "ts": datetime.now().strftime("%Y-%m-%d"),
                    })
            capture = True
            current_article = line + "\n"
        elif capture:
            if re.match(r'(الباب|الفصل)\s+', line) and len(current_article) > 100:
                clean = re.sub(r'\s+', ' ', current_article).strip()
                if sum(1 for c in clean if '\u0600' <= c <= '\u06ff') > 15:
                    records.append({
                        "text": clean[:1200],
                        "article": f"مادة {len(records)+1}",
                        "law_name": current_law,
                        "source": source,
                        "ts": datetime.now().strftime("%Y-%m-%d"),
                    })
                capture = False
                current_article = ""
            else:
                current_article += line + "\n"

    if capture and len(current_article) > 50:
        clean = re.sub(r'\s+', ' ', current_article).strip()
        if sum(1 for c in clean if '\u0600' <= c <= '\u06ff') > 15:
            records.append({
                "text": clean[:1200],
                "article": "مادة أخيرة",
                "law_name": current_law,
                "source": source,
                "ts": datetime.now().strftime("%Y-%m-%d"),
            })
    return records

def extract_laws_from_pdf(raw: bytes, source: str = "") -> List[Dict]:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for pg in pdf.pages:
                t = pg.extract_text() or ""
                if t.strip():
                    text += t + "\n"
    except Exception:
        try:
            import PyPDF2
            for pg in PyPDF2.PdfReader(io.BytesIO(raw)).pages:
                t = pg.extract_text() or ""
                if t.strip():
                    text += t + "\n"
        except Exception:
            pass
    return extract_laws_from_text(text, source) if text else []

def extract_laws_from_docx(raw: bytes, source: str = "") -> List[Dict]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return extract_laws_from_text(text, source)
    except Exception:
        return []

class DocIntel:
    def extract(self, f) -> str:
        ext = (getattr(f, "name", "") or "").rsplit(".", 1)[-1].lower()
        raw = _bytes(f)
        try:
            if ext == "pdf":
                return self._pdf(raw)
            if ext == "docx":
                return self._docx(raw)
            if ext in ("html", "htm"):
                return _norm(re.sub(r'<[^>]+>', '', raw.decode("utf-8", errors="ignore")))
            if ext == "json":
                import json
                return _norm(json.dumps(json.loads(raw.decode("utf-8", errors="ignore")), ensure_ascii=False))
            if ext == "csv":
                import csv
                rows = list(csv.reader(io.StringIO(raw.decode("utf-8", errors="ignore"))))
                return _norm("\n".join(" | ".join(r) for r in rows))
            return _norm(raw.decode("utf-8", errors="ignore"))
        except Exception:
            return ""

    def _pdf(self, raw: bytes) -> str:
        parts = []
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for pg in pdf.pages:
                    t = pg.extract_text() or ""
                    if t.strip():
                        parts.append(t)
        except Exception:
            try:
                import PyPDF2
                for pg in PyPDF2.PdfReader(io.BytesIO(raw)).pages:
                    t = pg.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            except Exception:
                pass
        return _norm("\n".join(parts))

    def _docx(self, raw: bytes) -> str:
        try:
            from docx import Document
            return _norm("\n".join(p.text for p in Document(io.BytesIO(raw)).paragraphs if p.text))
        except Exception:
            return ""

    def entities(self, t: str) -> dict:
        base_entities = {
            "parties": list(set(re.findall(r"(?:المدعي|المدعى عليه|الشركة|المؤسسة|الموظف|الهيئة)", t or ""))),
            "amounts": re.findall(r"[\d,]+\s*(?:ريال|درهم|دولار)", t or ""),
            "articles": re.findall(r"المادة\s*[\u0600-\u06FF\d]+", t or ""),
            "dates": re.findall(r"\d{1,2}/\d{1,2}/\d{2,4}", t or ""),
        }
        labor_entities = extract_labor_entities(t)
        base_entities.update(labor_entities)
        return base_entities
